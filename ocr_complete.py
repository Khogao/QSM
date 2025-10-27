#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Quicord - Quick OCR Documents
Vietnamese Document Intelligence with AI
"""

import sys
import io
import os

# Force UTF-8 encoding everywhere (fix Vietnamese character issues)
if sys.platform == "win32":
    # Windows: Fix console encoding
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8")
    # Set environment variable for subprocess
    os.environ["PYTHONIOENCODING"] = "utf-8"
else:
    # Mac/Linux: Ensure UTF-8
    import locale
    locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')

# Set default file encoding to UTF-8
import _locale
_locale._getdefaultlocale = (lambda *args: ['en_US', 'UTF-8'])
from pathlib import Path
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage, Table as RLTable, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pypdf import PdfReader, PdfWriter
from ebooklib import epub
from docx.table import Table as DocxTable
import time
import re
import uuid
import cv2
import numpy as np

# Docling IBM Models for advanced figure detection
try:
    from docling_ibm_models.document_figures.document_figure_classifier_predictor import DocumentFigureClassifierPredictor
    DOCLING_FIGURES_AVAILABLE = True
except ImportError:
    DOCLING_FIGURES_AVAILABLE = False
    print("[⚠] Warning: docling-ibm-models not installed. Advanced figure detection disabled.")
    print("    Install: pip install 'docling-ibm-models[opencv-python]>=3.10.1'")

# Supported formats - now includes PDF!
IMAGE_FORMATS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff"}
PDF_FORMAT = {".pdf"}
ALL_FORMATS = IMAGE_FORMATS | PDF_FORMAT

def header():
    print("="*70)
    print("� Quicord v3.0 - Quick OCR Documents")
    print("="*70)
    print("🇻🇳 Vietnamese Document Intelligence with AI")
    print("="*70)
    print("✨ v3.0 Features:")
    print("   • 16 figure types (QR, barcode, signature, stamp, charts)")
    print("   • Auto document classification (invoice, contract, blueprint...)")
    print("   • Smart filename suggestions")
    print("   • AI-powered text restructuring (coming soon)")
    print("="*70)
    print("📄 Formats: JPG, PNG, BMP, TIFF, PDF (scanned)")
    print("🌍 Languages: Vietnamese + English (95-98% accuracy)")
    print("📊 Export: Word, PDF, Excel, Markdown, EPUB, Text")
    print("🖥️  Platforms: Windows 11 + Mac OS")
    print("="*70)

def setup():
    """Initialize OCR converter for images and PDFs"""
    print("🚀 Starting Quicord OCR engine...")
    
    # OCR options
    ocr_opts = EasyOcrOptions(
        lang=["vi", "en"],
        force_full_page_ocr=True
    )
    
    # Pipeline for PDF
    pdf_pipeline_opts = PdfPipelineOptions()
    pdf_pipeline_opts.do_ocr = True
    pdf_pipeline_opts.do_table_structure = True
    pdf_pipeline_opts.ocr_options = ocr_opts
    
    # Pipeline for images
    img_pipeline_opts = PdfPipelineOptions()
    img_pipeline_opts.do_ocr = True
    img_pipeline_opts.do_table_structure = True
    img_pipeline_opts.ocr_options = ocr_opts
    
    # Create converter with configuration for both types
    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_pipeline_opts),
            InputFormat.IMAGE: PdfFormatOption(pipeline_options=img_pipeline_opts)
        }
    )
    
    print("✓ OCR engine ready (Vietnamese + English)")
    print("✓ PDF scan support")
    print("✓ Image support")
    print("✓ v3.0: 16 figure types detection (QR, barcode, signature, stamp, charts...)")
    print("✓ v3.0: Document type auto-detection (invoice, contract, blueprint...)")
    print("✓ v3.0: Smart filename suggestions")
    print("✓ Cross-platform: Windows 11 + Mac OS")
    print()
    return converter

def detect_document_figures(image_path):
    """
    Detect document figures using Docling IBM Models (v3.0)
    
    Detects 16 types: QR code, barcode, signature, stamp, charts (bar, line, pie, flow), 
    chemistry structures, icons, logos, maps, screenshots, remote sensing images
    
    Args:
        image_path: Path to image file (or numpy array from PDF page)
    
    Returns:
        List of figure dicts: [{"type": "qr_code", "confidence": 0.95, "rect": (x,y,w,h)}, ...]
    """
    if not DOCLING_FIGURES_AVAILABLE:
        # Fallback: No figure detection if docling-ibm-models not installed
        return []
    
    try:
        # Initialize classifier (singleton pattern - only load once)
        if not hasattr(detect_document_figures, '_predictor'):
            detect_document_figures._predictor = DocumentFigureClassifierPredictor()
        
        predictor = detect_document_figures._predictor
        
        # Read image
        if isinstance(image_path, (str, Path)):
            img = cv2.imread(str(image_path))
        else:
            # Already numpy array (from PDF page)
            img = image_path
        
        if img is None:
            return []
        
        # Convert to PIL Image for predictor
        from PIL import Image
        if isinstance(img, np.ndarray):
            # OpenCV uses BGR, PIL uses RGB
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(img_rgb)
        else:
            pil_img = Image.open(str(image_path))
        
        # Run prediction
        predictions = predictor.predict([pil_img])
        
        results = []
        for pred_list in predictions:
            for class_name, confidence in pred_list:
                # Filter low confidence (< 0.5)
                if confidence < 0.5:
                    continue
                
                results.append({
                    'type': class_name,  # e.g. "qr_code", "signature", "stamp", "bar_code"
                    'confidence': confidence,
                    'rect': None,  # DocumentFigureClassifier doesn't provide bounding boxes
                })
        
        return results
        
    except Exception as e:
        print(f"    [⚠] Figure detection error: {e}")
        return []

def detect_document_type(ocr_text, figures_detected):
    """
    Auto-detect document type based on OCR text and detected figures (v3.0)
    
    Document types:
    - invoice (hóa đơn): Has VAT, tax code, QR, signature, stamp
    - contract (hợp đồng): Has parties, signature, stamp
    - blueprint (bản vẽ): Has charts, scale indicators
    - certificate (chứng chỉ): Has signature, stamp, certification keywords
    - receipt (biên lai): Similar to invoice but simpler
    - other: Unknown/generic document
    
    Args:
        ocr_text: Extracted text from OCR
        figures_detected: List of detected figures from detect_document_figures()
    
    Returns:
        (doc_type: str, confidence: float, keywords_found: list)
    """
    text_lower = ocr_text.lower()
    
    # Extract figure types
    figure_types = {fig['type'] for fig in figures_detected}
    
    # Define detection rules
    scores = {
        'invoice': 0.0,
        'contract': 0.0,
        'blueprint': 0.0,
        'certificate': 0.0,
        'receipt': 0.0,
        'other': 0.0
    }
    
    keywords_found = []
    
    # Invoice detection - IMPROVED with more Vietnamese variants
    invoice_keywords = [
        # Main invoice terms
        ('hóa đơn', 0.35), ('hoá đơn', 0.35), ('hoa don', 0.35), ('invoice', 0.35),
        ('hóa đơn gtgt', 0.4), ('hóa đơn vat', 0.4),
        # Tax terms
        ('vat', 0.25), ('gtgt', 0.25), ('thuế', 0.15), ('thue', 0.15),
        ('mst', 0.2), ('mã số thuế', 0.2), ('ma so thue', 0.2), ('tax code', 0.2),
        # Payment terms
        ('thành tiền', 0.12), ('thanh tien', 0.12),
        ('tổng cộng', 0.12), ('tong cong', 0.12), ('total', 0.12),
        ('tổng tiền', 0.12), ('tong tien', 0.12), ('sum', 0.12),
        # VAT specific
        ('tiền thuế', 0.1), ('tien thue', 0.1),
        ('cộng tiền hàng', 0.1), ('cong tien hang', 0.1),
    ]
    for keyword, weight in invoice_keywords:
        if keyword in text_lower:
            scores['invoice'] += weight
            if keyword not in keywords_found:  # Avoid duplicates
                keywords_found.append(keyword)
    
    if 'qr_code' in figure_types:
        scores['invoice'] += 0.25
        keywords_found.append('QR code')
    if 'signature' in figure_types:
        scores['invoice'] += 0.15
        keywords_found.append('signature')
    if 'stamp' in figure_types:
        scores['invoice'] += 0.15
        keywords_found.append('stamp')
    if 'bar_code' in figure_types:
        scores['invoice'] += 0.1
        keywords_found.append('barcode')
    
    # Contract detection - MUCH IMPROVED based on Vietnamese legal contracts
    contract_keywords = [
        # Main contract terms
        ('hợp đồng', 0.45), ('hop dong', 0.45), ('hopdong', 0.45), ('contract', 0.45),
        ('hợp đồng thế chấp', 0.5), ('hop dong the chap', 0.5),  # Mortgage contract
        ('hợp đồng tín dụng', 0.5), ('hop dong tin dung', 0.5),  # Credit contract
        # Parties (very strong indicators!)
        ('bên a', 0.3), ('ben a', 0.3), ('party a', 0.3),
        ('bên b', 0.3), ('ben b', 0.3), ('party b', 0.3),
        ('bên c', 0.25), ('ben c', 0.25), ('party c', 0.25),
        ('bên thế chấp', 0.3), ('ben the chap', 0.3),  # Mortgagor
        ('bên nhận thế chấp', 0.3), ('ben nhan the chap', 0.3),  # Mortgagee
        ('bên vay', 0.25), ('ben vay', 0.25),  # Borrower
        ('bên cho vay', 0.25), ('ben cho vay', 0.25),  # Lender
        ('người đại diện', 0.2), ('nguoi dai dien', 0.2),  # Representative
        # Legal terms (super strong!)
        ('điều khoản', 0.2), ('dieu khoan', 0.2), ('terms', 0.2), ('clause', 0.2),
        ('điều 1', 0.15), ('điều 2', 0.15), ('dieu 1', 0.15), ('dieu 2', 0.15),  # Articles
        ('quyền và nghĩa vụ', 0.25), ('quyen va nghia vu', 0.25),  # Rights and obligations
        ('thỏa thuận', 0.15), ('thoa thuan', 0.15), ('agreement', 0.15),
        ('cam kết', 0.15), ('cam ket', 0.15), ('commitment', 0.15),
        # Signatures
        ('chữ ký', 0.12), ('chu ky', 0.12), ('signature', 0.12),
        ('ký kết', 0.15), ('ky ket', 0.15), ('signed', 0.15),
        ('đại diện pháp luật', 0.2), ('dai dien phap luat', 0.2),  # Legal representative
        # Common contract phrases
        ('cơ sở pháp lý', 0.1), ('co so phap ly', 0.1),  # Legal basis
        ('phạm vi áp dụng', 0.1), ('pham vi ap dung', 0.1),  # Scope
        ('hiệu lực', 0.1), ('hieu luc', 0.1),  # Effectiveness
    ]
    for keyword, weight in contract_keywords:
        if keyword in text_lower:
            scores['contract'] += weight
            if keyword not in keywords_found:
                keywords_found.append(keyword)
    
    if 'signature' in figure_types:
        scores['contract'] += 0.3
    if 'stamp' in figure_types:
        scores['contract'] += 0.2
    
    # Blueprint detection
    blueprint_keywords = [
        ('bản vẽ', 0.4), ('ban ve', 0.4), ('blueprint', 0.4),
        ('tỷ lệ', 0.2), ('scale', 0.2), ('ty le', 0.2),
        ('kích thước', 0.15), ('dimension', 0.15),
        ('mặt cắt', 0.1), ('section', 0.1)
    ]
    for keyword, weight in blueprint_keywords:
        if keyword in text_lower:
            scores['blueprint'] += weight
            keywords_found.append(keyword)
    
    chart_types = {'bar_chart', 'line_chart', 'pie_chart', 'flow_chart'}
    if chart_types & figure_types:
        scores['blueprint'] += 0.3
        keywords_found.append('charts')
    
    # Certificate detection
    cert_keywords = [
        ('chứng nhận', 0.4), ('chung nhan', 0.4), ('certificate', 0.4),
        ('giấy chứng nhận', 0.4), ('certification', 0.4),
        ('cấp cho', 0.15), ('issued to', 0.15),
        ('có giá trị', 0.1), ('valid until', 0.1)
    ]
    for keyword, weight in cert_keywords:
        if keyword in text_lower:
            scores['certificate'] += weight
            keywords_found.append(keyword)
    
    if 'signature' in figure_types:
        scores['certificate'] += 0.25
    if 'stamp' in figure_types:
        scores['certificate'] += 0.25
    
    # Receipt detection
    receipt_keywords = [
        ('biên lai', 0.4), ('bien lai', 0.4), ('receipt', 0.4),
        ('phiếu thu', 0.3), ('phieu thu', 0.3),
        ('đã nhận', 0.15), ('da nhan', 0.15), ('received', 0.15)
    ]
    for keyword, weight in receipt_keywords:
        if keyword in text_lower:
            scores['receipt'] += weight
            keywords_found.append(keyword)
    
    # Find best match
    doc_type = max(scores, key=scores.get)
    confidence = min(scores[doc_type], 1.0)  # Cap at 1.0
    
    # If confidence too low, mark as "other"
    if confidence < 0.3:
        doc_type = 'other'
        confidence = 0.5
    
    return doc_type, confidence, keywords_found

def suggest_filename(doc_type, ocr_text, figures_detected, original_filename):
    """
    Smart filename suggestions based on document type (v3.0)
    
    Args:
        doc_type: Document type from detect_document_type()
        ocr_text: Extracted OCR text
        figures_detected: List of detected figures
        original_filename: Original file name (fallback)
    
    Returns:
        Suggested filename (sanitized, ready to use)
    """
    import re
    from datetime import datetime
    
    # Extract QR data if available
    qr_data = None
    for fig in figures_detected:
        if fig['type'] == 'qr_code' and 'data' in fig:
            qr_data = fig.get('data', '')[:50]  # First 50 chars
            break
    
    # Sanitize function
    def sanitize(text):
        # Remove special chars, keep Vietnamese, alphanumeric, spaces, dashes
        text = re.sub(r'[^\w\s\-àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]', '', text)
        text = re.sub(r'\s+', '_', text)  # Replace spaces with underscores
        text = text.strip('_')
        return text[:50]  # Max 50 chars
    
    # Extract date (YYYY-MM-DD, DD/MM/YYYY, etc.)
    date_match = re.search(r'(\d{4}[-/\.]\d{1,2}[-/\.]\d{1,2})|(\d{1,2}[-/\.]\d{1,2}[-/\.]\d{4})', ocr_text)
    date_str = date_match.group(0).replace('/', '-').replace('.', '-') if date_match else datetime.now().strftime('%Y-%m-%d')
    
    suggested = None
    
    if doc_type == 'invoice':
        # Extract: invoice number, company name
        invoice_num = re.search(r'(số|so|no\.?|#)\s*[:.]?\s*(\w+[-/]?\d+)', ocr_text.lower())
        company = re.search(r'(công ty|cong ty|company)\s+(.{5,30})', ocr_text.lower())
        
        parts = ['Invoice']
        if invoice_num:
            parts.append(sanitize(invoice_num.group(2)))
        if company:
            parts.append(sanitize(company.group(2)))
        parts.append(date_str)
        
        suggested = '_'.join(parts)
    
    elif doc_type == 'contract':
        # Extract: contract number, parties
        contract_num = re.search(r'(hợp đồng|hop dong|contract)\s+số\s*[:.]?\s*(\w+[-/]?\d+)', ocr_text.lower())
        party_a = re.search(r'bên a\s*[:.]?\s*(.{5,30})', ocr_text.lower())
        
        parts = ['Contract']
        if contract_num:
            parts.append(sanitize(contract_num.group(2)))
        if party_a:
            parts.append(sanitize(party_a.group(1)))
        parts.append(date_str)
        
        suggested = '_'.join(parts)
    
    elif doc_type == 'blueprint':
        # Extract: project name, scale, sheet number
        project = re.search(r'(dự án|du an|project)\s*[:.]?\s*(.{5,30})', ocr_text.lower())
        scale = re.search(r'(tỷ lệ|ty le|scale)\s*[:.]?\s*([\d:]+)', ocr_text.lower())
        sheet = re.search(r'(tờ|sheet|page)\s*[:.]?\s*(\d+)', ocr_text.lower())
        
        parts = ['Blueprint']
        if project:
            parts.append(sanitize(project.group(2)))
        if scale:
            parts.append(f"Scale_{scale.group(2)}")
        if sheet:
            parts.append(f"Sheet_{sheet.group(2)}")
        
        suggested = '_'.join(parts)
    
    elif doc_type == 'certificate':
        # Extract: certificate type, recipient
        cert_type = re.search(r'(chứng nhận|chung nhan|certificate)\s+(.{5,30})', ocr_text.lower())
        recipient = re.search(r'(cấp cho|cap cho|issued to)\s*[:.]?\s*(.{5,30})', ocr_text.lower())
        
        parts = ['Certificate']
        if cert_type:
            parts.append(sanitize(cert_type.group(2)))
        if recipient:
            parts.append(sanitize(recipient.group(1)))
        parts.append(date_str)
        
        suggested = '_'.join(parts)
    
    elif doc_type == 'receipt':
        # Extract: receipt number
        receipt_num = re.search(r'(biên lai|bien lai|receipt)\s+số\s*[:.]?\s*(\w+[-/]?\d+)', ocr_text.lower())
        
        parts = ['Receipt']
        if receipt_num:
            parts.append(sanitize(receipt_num.group(2)))
        parts.append(date_str)
        
        suggested = '_'.join(parts)
    
    else:  # other
        # Use original filename + date
        original_clean = sanitize(Path(original_filename).stem)
        suggested = f"{original_clean}_{date_str}"
    
    # Fallback if extraction failed
    if not suggested or suggested == doc_type:
        suggested = f"{doc_type}_{date_str}_{sanitize(original_filename)}"
    
    return suggested

def is_pdf_scanned(pdf_path):
    """Check if PDF is scanned (image-based) or has text"""
    try:
        reader = PdfReader(pdf_path)
        total_text = ""
        
        # Check first 3 pages
        for page_num in range(min(3, len(reader.pages))):
            page = reader.pages[page_num]
            text = page.extract_text()
            total_text += text
        
        # If very little text, likely scanned
        if len(total_text.strip()) < 50:
            return True, "scanned"
        else:
            return False, "text-based"
    except:
        return True, "unknown"

def proc(p, c, o):
    """Process single file (image or PDF) and return OCR result"""
    try:
        p = Path(p).resolve()
        if not p.exists(): 
            print(f"[✗] Not found: {p}")
            return None
        
        suffix = p.suffix.lower()
        if suffix not in ALL_FORMATS:
            print(f"[!] Unsupported format: {suffix}")
            return None
        
        # File info
        file_size_mb = p.stat().st_size / (1024 * 1024)
        print(f"\n[→] {p.name} ({file_size_mb:.2f} MB)")
        
        # Check if PDF is scanned
        if suffix == ".pdf":
            is_scanned, pdf_type = is_pdf_scanned(p)
            print(f"    Type: PDF ({pdf_type})")
            
            if not is_scanned:
                print(f"    [!] Warning: This PDF already has text.")
                print(f"        OCR will still run but may not be necessary.")
        else:
            print(f"    Type: Image")
        
        # Detect document figures (QR, barcode, signature, stamp, charts) - v3.0
        print("    🔍 Đang phát hiện figures (QR, signature, stamp, charts)...")
        figures = []
        try:
            figures = detect_document_figures(p)
        except Exception as e:
            print(f"    [⚠] Figure detection bỏ qua: {e}")
        
        if figures:
            print(f"    [✓] Tìm thấy {len(figures)} figure(s)")
            for idx, fig in enumerate(figures, 1):
                print(f"        {idx}. {fig['type']} (confidence: {fig['confidence']:.2f})")
        
        # Run OCR
        print("    Đang xử lý OCR...")
        t = time.time()
        result = c.convert(str(p))
        duration = time.time() - t
        
        # Extract text and tables
        txt = result.document.export_to_markdown()
        
        if not txt.strip():
            print("    [✗] Không trích xuất được văn bản")
            return None
        
        # Analyze tables (Docling detects tables automatically)
        tables_found = txt.count('|')  # Markdown tables use |
        has_tables = tables_found > 10  # More than 10 pipes = likely has tables
        
        word_count = len(txt.split())
        print(f"    [✓] Xong trong {duration:.1f}s")
        print(f"    [✓] Trích xuất {word_count} từ")
        if has_tables:
            print(f"    [✓] Phát hiện bảng biểu (Docling table detection)")
        
        # v3.0: Auto-detect document type
        doc_type, doc_confidence, doc_keywords = detect_document_type(txt, figures)
        print(f"    [✓] Document type: {doc_type.upper()} (confidence: {doc_confidence:.2f})")
        if doc_keywords:
            print(f"        Keywords: {', '.join(doc_keywords[:5])}")
        
        # v3.0: Suggest filename
        suggested_name = suggest_filename(doc_type, txt, figures, p.name)
        print(f"    [💡] Suggested filename: {suggested_name}")
        
        # Return result dict with enhanced metadata (v3.0)
        return {
            'path': p,
            'text': txt,
            'word_count': word_count,
            'time': duration,
            'type': 'pdf' if suffix == '.pdf' else 'image',
            'has_tables': has_tables,
            'figures': figures,  # v3.0: Detected figures (QR, signature, stamp, charts...)
            'document_type': doc_type,  # v3.0: Auto-detected document type
            'document_confidence': doc_confidence,  # v3.0: Document type confidence
            'document_keywords': doc_keywords,  # v3.0: Keywords found
            'suggested_filename': suggested_name,  # v3.0: Smart filename suggestion
            'raw_result': result  # Keep raw result for advanced processing
        }
        
    except Exception as e:
        print(f"    [✗] ERROR: {e}")
        import traceback
        traceback.print_exc()
        return None

def extract_page_number(filename):
    """Extract page number from filename for sorting"""
    # Try to find numbers in filename
    numbers = re.findall(r'\d+', filename)
    if numbers:
        # Return first number found (usually page number)
        return int(numbers[0])
    return 0

def sort_files_by_page(files):
    """Sort files by page number extracted from filename"""
    return sorted(files, key=lambda f: (extract_page_number(f.stem), f.name))

def create_merged_document(results, output_dir, doc_name="merged_document"):
    """Create merged Word and Markdown from OCR results"""
    
    # Sort results by page number
    sorted_results = sorted(results, key=lambda r: extract_page_number(r['path'].stem))
    
    print("\n[*] Creating merged document...")
    print(f"    Page order: {', '.join([r['path'].stem[:20] + '...' if len(r['path'].stem) > 20 else r['path'].stem for r in sorted_results])}")
    
    # Create merged Markdown
    md_content = f"# Merged OCR Document\n\n"
    md_content += f"*Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
    md_content += f"*Total pages: {len(sorted_results)}*\n\n"
    md_content += "---\n\n"
    
    for idx, result in enumerate(sorted_results, 1):
        md_content += f"## Page {idx} - {result['path'].name}\n\n"
        md_content += result['text']
        md_content += f"\n\n---\n\n"
    
    md_file = output_dir / f"{doc_name}.md"
    md_file.write_text(md_content, encoding="utf-8")
    print(f"    [✓] Markdown: {md_file.name}")
    
    # Create merged Word document
    doc = Document()
    
    # Title page
    title = doc.add_heading('OCR Merged Document', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"Total Pages: {len(sorted_results)}\n").italic = True
    meta.add_run(f"Total Words: {sum(r['word_count'] for r in sorted_results)}\n").italic = True
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Add each page
    for idx, result in enumerate(sorted_results, 1):
        # Page header
        page_heading = doc.add_heading(f'Trang {idx}', 1)
        
        # Source info
        source = doc.add_paragraph()
        source.add_run(f"Nguồn: {result['path'].name}\n").italic = True
        source.add_run(f"Số từ: {result['word_count']}").italic = True
        if result.get('has_tables'):
            source.add_run("\n✓ Có bảng biểu").bold = True
        
        # Add separator
        doc.add_paragraph("_" * 50)
        
        # Content - improved table handling
        text_content = result['text']
        
        # Split by double newline to preserve structure
        paragraphs = text_content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Check if this is a Markdown table
            if '|' in para_text and para_text.count('|') > 2:
                # This is likely a table
                try:
                    lines = para_text.strip().split('\n')
                    # Filter out separator lines (like |---|---|)
                    table_lines = [line for line in lines if not all(c in '|-: ' for c in line)]
                    
                    if len(table_lines) > 0:
                        # Parse table
                        rows = []
                        for line in table_lines:
                            # Split by | and clean
                            cells = [cell.strip() for cell in line.split('|')]
                            # Remove empty first/last cells (from leading/trailing |)
                            cells = [c for c in cells if c]
                            if cells:
                                rows.append(cells)
                        
                        if rows:
                            # Create Word table
                            num_cols = max(len(row) for row in rows)
                            table = doc.add_table(rows=len(rows), cols=num_cols)
                            table.style = 'Light Grid Accent 1'
                            
                            # Fill table
                            for i, row_data in enumerate(rows):
                                for j, cell_text in enumerate(row_data):
                                    if j < num_cols:
                                        cell = table.rows[i].cells[j]
                                        cell.text = cell_text
                                        # Bold first row (header)
                                        if i == 0:
                                            cell.paragraphs[0].runs[0].font.bold = True
                            
                            doc.add_paragraph()  # Space after table
                            continue
                except Exception as e:
                    # If table parsing fails, fall back to normal text
                    print(f"    [⚠] Lỗi parse bảng: {e}")
            
            # Normal text (not a table)
            if para_text.startswith('#'):
                # Heading
                level = min(para_text.count('#', 0, 3), 2) + 1
                text_clean = para_text.lstrip('#').strip()
                doc.add_heading(text_clean, level)
            else:
                # Normal paragraph
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.line_spacing = 1.5
        
        # Page break (except last page)
        if idx < len(sorted_results):
            doc.add_page_break()
    
    docx_file = output_dir / f"{doc_name}.docx"
    doc.save(str(docx_file))
    print(f"    [✓] Word: {docx_file.name}")
    
    return md_file, docx_file

def create_pdf_from_images(image_paths, output_path):
    """Create PDF from original images (preserving quality)"""
    try:
        print(f"\n[*] Creating PDF from images...")
        
        # Filter only image files
        valid_images = [p for p in image_paths if p.suffix.lower() in IMAGE_FORMATS]
        
        if not valid_images:
            print("    [!] No image files to convert")
            return None
        
        # Open all images and convert to RGB
        images = []
        for img_path in valid_images:
            img = Image.open(img_path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)
        
        # Save as PDF
        if images:
            images[0].save(
                output_path,
                save_all=True,
                append_images=images[1:],
                resolution=100.0,
                quality=95,
                optimize=True
            )
            print(f"    [✓] Image PDF: {output_path.name} ({len(images)} pages)")
            return output_path
    except Exception as e:
        print(f"    [✗] Failed to create image PDF: {e}")
        return None

def create_pdf_from_ocr(results, output_path):
    """Create PDF from OCR text content with Vietnamese font support"""
    try:
        print(f"\n[*] Tạo PDF từ văn bản OCR...")
        
        # Register Vietnamese-compatible font (DejaVu Sans has Vietnamese)
        # Try to use system fonts that support Vietnamese
        try:
            # Windows fonts
            pdfmetrics.registerFont(TTFont('VietnameseFont', 'C:/Windows/Fonts/arial.ttf'))
            pdfmetrics.registerFont(TTFont('VietnameseBold', 'C:/Windows/Fonts/arialbd.ttf'))
            font_name = 'VietnameseFont'
            font_bold = 'VietnameseBold'
            print("    [✓] Đã load font Arial (hỗ trợ tiếng Việt)")
        except:
            try:
                # Fallback to DejaVu (included in reportlab)
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                font_name = 'STSong-Light'
                font_bold = 'STSong-Light'
                print("    [!] Dùng font dự phòng (có thể không đẹp)")
            except:
                font_name = 'Helvetica'
                font_bold = 'Helvetica-Bold'
                print("    [⚠] Cảnh báo: Không tìm thấy font tiếng Việt, dấu có thể bị lỗi")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # Styles with Vietnamese font
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_bold,
            fontSize=24,
            textColor='#2C3E50',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=font_bold,
            fontSize=16,
            textColor='#34495E',
            spaceAfter=12,
            spaceBefore=12
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=10
        )
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            textColor='#7F8C8D',
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        # Build content
        story = []
        
        # Title page
        story.append(Paragraph("OCR Merged Document", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
        story.append(Paragraph(f"Total Pages: {len(results)}", meta_style))
        story.append(Paragraph(f"Total Words: {sum(r['word_count'] for r in results)}", meta_style))
        story.append(PageBreak())
        
        # Add each page
        for idx, result in enumerate(results, 1):
            # Page header
            story.append(Paragraph(f"Page {idx}", heading_style))
            story.append(Paragraph(f"Source: {result['path'].name}", meta_style))
            
            # QR codes section
            if result.get('qr_codes'):
                story.append(Spacer(1, 0.05*inch))
                qr_text = f"<b>QR Codes:</b> Found {len(result['qr_codes'])}"
                story.append(Paragraph(qr_text, meta_style))
                for qr_idx, qr in enumerate(result['qr_codes'], 1):
                    qr_data = qr['data'][:100]  # Limit length
                    story.append(Paragraph(f"  QR{qr_idx}: {qr_data}", meta_style))
            
            story.append(Spacer(1, 0.1*inch))
            
            # Content with improved table handling
            text = result['text']
            paragraphs = text.split('\n\n')
            
            for para in paragraphs:
                if not para.strip():
                    continue
                
                # Check if this is a Markdown table
                if '|' in para and para.count('|') > 2:
                    # This is likely a table
                    try:
                        lines = para.strip().split('\n')
                        # Filter out separator lines (like |---|---|)
                        table_lines = [line for line in lines if not all(c in '|-: ' for c in line)]
                        
                        if len(table_lines) > 0:
                            # Parse table
                            rows = []
                            for line in table_lines:
                                # Split by | and clean
                                cells = [cell.strip() for cell in line.split('|')]
                                # Remove empty first/last cells (from leading/trailing |)
                                cells = [c for c in cells if c]
                                if cells:
                                    # Wrap each cell in Paragraph for Vietnamese font support
                                    cell_paras = [Paragraph(cell, normal_style) for cell in cells]
                                    rows.append(cell_paras)
                            
                            if rows:
                                # Create ReportLab Table
                                table = RLTable(rows)
                                
                                # Table styling
                                table_style = TableStyle([
                                    # Header row (first row)
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                    ('FONTNAME', (0, 0), (-1, 0), font_bold),
                                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                                    ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                                    
                                    # All cells
                                    ('FONTNAME', (0, 1), (-1, -1), font_name),
                                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                                    ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                                    
                                    # Borders
                                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                    ('BOX', (0, 0), (-1, -1), 1, colors.black),
                                    
                                    # Padding
                                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                ])
                                table.setStyle(table_style)
                                
                                story.append(table)
                                story.append(Spacer(1, 0.1*inch))
                                continue
                    except Exception as e:
                        # If table parsing fails, fall back to normal text
                        print(f"    [⚠] Lỗi parse bảng cho PDF: {e}")
                
                # Normal text (not a table)
                clean_text = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Check if heading
                if clean_text.startswith('#'):
                    clean_text = clean_text.lstrip('#').strip()
                    story.append(Paragraph(clean_text, heading_style))
                else:
                    story.append(Paragraph(clean_text, normal_style))
            
            # Page break (except last)
            if idx < len(results):
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        print(f"    [✓] Text PDF: {output_path.name} (with table formatting)")
        return output_path
        
    except Exception as e:
        print(f"    [✗] Failed to create text PDF: {e}")
        import traceback
        traceback.print_exc()
        return None

def extract_tables_to_excel(results, output_path):
    """
    [PHASE 2 FEATURE] Extract tables from OCR results to Excel
    
    This function extracts all tables detected by Docling and exports them
    to a formatted Excel file. Perfect for Vietnamese invoice OCR!
    
    Args:
        results: List of OCR result dicts
        output_path: Path to save Excel file
    
    Returns:
        Path to Excel file or None if failed
    """
    try:
        # Check if openpyxl is installed (Phase 2 dependency)
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        except ImportError:
            print("    [!] Excel export cần cài openpyxl:")
            print("    pip install openpyxl")
            return None
        
        print(f"\n[*] Trích xuất bảng biểu ra Excel...")
        
        # Create workbook
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        table_count = 0
        
        # Process each result
        for idx, result in enumerate(results, 1):
            text = result['text']
            
            # Find all tables in markdown
            paragraphs = text.split('\n\n')
            page_table_count = 0
            
            for para in paragraphs:
                if '|' not in para or para.count('|') < 3:
                    continue
                
                # Parse table
                lines = para.strip().split('\n')
                table_lines = [line for line in lines if not all(c in '|-: ' for c in line)]
                
                if len(table_lines) < 2:  # Need at least header + 1 row
                    continue
                
                # Parse rows
                rows = []
                for line in table_lines:
                    cells = [cell.strip() for cell in line.split('|')]
                    cells = [c for c in cells if c]
                    if cells:
                        rows.append(cells)
                
                if not rows:
                    continue
                
                # Create worksheet for this table
                page_table_count += 1
                table_count += 1
                sheet_name = f"Trang{idx}_Bảng{page_table_count}"[:31]  # Excel limit
                ws = wb.create_sheet(title=sheet_name)
                
                # Write data
                for row_idx, row_data in enumerate(rows, 1):
                    for col_idx, cell_value in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=cell_value)
                        
                        # Format header row
                        if row_idx == 1:
                            cell.font = Font(bold=True, color="FFFFFF")
                            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                        else:
                            cell.alignment = Alignment(horizontal='left', vertical='center')
                        
                        # Add borders
                        thin_border = Border(
                            left=Side(style='thin'),
                            right=Side(style='thin'),
                            top=Side(style='thin'),
                            bottom=Side(style='thin')
                        )
                        cell.border = thin_border
                
                # Auto-adjust column width
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)  # Max 50 chars
                    ws.column_dimensions[column_letter].width = adjusted_width
        
        if table_count == 0:
            print("    [!] Không tìm thấy bảng biểu nào")
            return None
        
        # Save workbook
        wb.save(output_path)
        print(f"    [✓] Excel: {output_path.name} ({table_count} bảng)")
        return output_path
        
    except Exception as e:
        print(f"    [✗] Lỗi xuất Excel: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_epub(results, output_path, book_title="OCR Document", author="QSM OCR"):
    """Create EPUB ebook from OCR results"""
    try:
        print(f"\n[*] Tạo EPUB ebook...")
        
        # Create book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(str(uuid.uuid4()))
        book.set_title(book_title)
        book.set_language('vi')  # Vietnamese
        book.add_author(author)
        
        # Create chapters
        chapters = []
        spine = ['nav']
        
        for idx, result in enumerate(results, 1):
            # Create chapter
            chapter = epub.EpubHtml(
                title=f'Page {idx}',
                file_name=f'page_{idx}.xhtml',
                lang='vi'
            )
            
            # Build HTML content
            content = f'<h1>Page {idx}</h1>'
            content += f'<p><em>Source: {result["path"].name}</em></p>'
            content += '<hr/>'
            
            # Convert markdown to HTML
            text = result['text']
            paragraphs = text.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Clean HTML
                    clean = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    if clean.startswith('#'):
                        # Heading
                        level = min(clean.count('#', 0, 3), 3)
                        text_clean = clean.lstrip('#').strip()
                        content += f'<h{level + 1}>{text_clean}</h{level + 1}>'
                    else:
                        # Paragraph
                        content += f'<p>{clean}</p>'
            
            chapter.content = content
            
            # Add to book
            book.add_item(chapter)
            chapters.append(chapter)
            spine.append(chapter)
        
        # Add default NCX and Nav files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = '''
            @namespace epub "http://www.idpf.org/2007/ops";
            body { font-family: Georgia, serif; margin: 2em; }
            h1 { color: #2C3E50; border-bottom: 2px solid #3498DB; padding-bottom: 0.3em; }
            h2, h3 { color: #34495E; }
            p { text-align: justify; line-height: 1.6; }
            em { color: #7F8C8D; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Create spine
        book.spine = spine
        
        # Add table of contents
        book.toc = tuple(chapters)
        
        # Write EPUB
        epub.write_epub(str(output_path), book, {})
        print(f"    [✓] EPUB: {output_path.name} ({len(chapters)} chapters)")
        return output_path
        
    except Exception as e:
        print(f"    [✗] Failed to create EPUB: {e}")
        import traceback
        traceback.print_exc()
        return None

def get_files():
    """Get files from user input (supports drag & drop)"""
    print("Cách thêm file:")
    print("  1. Kéo thả file/folder vào đây")
    print("  2. Gõ đường dẫn file")
    print("  3. Gõ 'done' khi xong\n")
    
    fs = []
    seen = set()  # Tránh trùng lặp
    
    while True:
        i = input(">>> ").strip()
        if not i: 
            continue
        if i.lower() in ["done", "q", "exit", "xong"]: 
            break
        
        # Bỏ PowerShell syntax: & '...' hoặc & "..."
        i = re.sub(r'^&\s+', '', i)
        
        # Xử lý Windows drag & drop với dấu ngoặc kép
        # Ví dụ: "C:\file1.pdf" "C:\file2.pdf" "C:\file3.pdf"
        # Hoặc PowerShell: 'C:\file1.pdf' 'C:\file2.pdf'
        paths = re.findall(r'"([^"]+)"', i)
        if not paths:
            paths = re.findall(r"'([^']+)'", i)
        if not paths:
            # Nếu không có dấu ngoặc, coi như 1 đường dẫn
            paths = [i.strip('"').strip("'")]
        
        for path_str in paths:
            try:
                # Windows path có thể có \ hoặc /, Python xử lý được cả 2
                p = Path(path_str.strip()).resolve()
                
                # Kiểm tra folder
                if p.is_dir():
                    print(f"   [→] Quét folder: {p.name}")
                    found_files = []
                    for fmt in ALL_FORMATS:
                        found_files.extend(p.glob(f"*{fmt}"))
                        found_files.extend(p.glob(f"*{fmt.upper()}"))
                    
                    for f in found_files:
                        if f not in seen:
                            fs.append(f)
                            seen.add(f)
                            print(f"      [+] {f.name}")
                    
                    if not found_files:
                        print(f"      [!] Không tìm thấy file nào")
                
                # Kiểm tra file
                elif p.exists() and p.is_file():
                    suffix = p.suffix.lower()
                    if suffix in ALL_FORMATS:
                        if p not in seen:
                            fs.append(p)
                            seen.add(p)
                            print(f"   [+] {p.name}")
                    else:
                        print(f"   [✗] Format không hỗ trợ: {p.name} (chỉ chấp nhận: {', '.join(ALL_FORMATS)})")
                else:
                    print(f"   [✗] Không tìm thấy file: {path_str}")
                    
            except Exception as e:
                print(f"   [✗] Lỗi đường dẫn: {path_str} ({e})")
    
    return fs

def main():
    header()
    converter = setup()
    
    # Lấy danh sách file
    files = get_files()
    
    if not files:
        print("\n[!] Không có file nào được chọn")
        input("Nhấn Enter để thoát...")
        return
    
    # Sắp xếp theo số trang
    files = sort_files_by_page(files)
    
    # Hiển thị danh sách
    print(f"\n{'='*70}")
    print(f"📄 Danh sách file cần xử lý ({len(files)} file):")
    for idx, f in enumerate(files, 1):
        print(f"  {idx}. {f.name}")
    print(f"{'='*70}\n")
    
    # Xác nhận
    print("⚡ Bắt đầu xử lý? (y/n): ", end='')
    confirm = input().strip().lower()
    if confirm not in ['y', 'yes', 'có', '']:
        print("❌ Đã hủy.")
        return
    
    # Tạo thư mục output
    output_dir = Path(__file__).parent / "ocr_output"
    output_dir.mkdir(exist_ok=True)
    
    # Xử lý tất cả file
    print(f"\n{'='*70}")
    print("⚙️  ĐANG XỬ LÝ...")
    print(f"{'='*70}")
    
    results = []
    start_time = time.time()
    
    for f in files:
        result = proc(f, converter, output_dir)
        if result:
            results.append(result)
    
    total_time = time.time() - start_time
    
    # Summary
    print(f"\n{'='*70}")
    print(f"TỔNG KẾT")
    print(f"{'='*70}")
    print(f"Đã xử lý: {len(results)}/{len(files)} file")
    print(f"Tổng thời gian: {total_time:.1f}s")
    print(f"Tổng số từ: {sum(r['word_count'] for r in results)}")
    
    # Count tables
    tables_count = sum(1 for r in results if r.get('has_tables', False))
    if tables_count > 0:
        print(f"📊 File có bảng biểu: {tables_count}")
    
    # Count QR codes
    qr_count = sum(len(r.get('qr_codes', [])) for r in results)
    if qr_count > 0:
        print(f"📱 Tìm thấy {qr_count} QR code(s)")
        # Show QR data
        for idx, r in enumerate(results, 1):
            if r.get('qr_codes'):
                print(f"   File {idx} ({r['path'].name}):")
                for qr_idx, qr in enumerate(r['qr_codes'], 1):
                    print(f"     QR{qr_idx}: {qr['data'][:60]}...")
    
    print(f"{'='*70}")
    
    if not results:
        print("\n[!] Không có kết quả OCR thành công")
        input("Nhấn Enter để thoát...")
        return
    
    # Create merged document
    print("\nTạo tài liệu gộp? (y/n): ", end='')
    choice = input().strip().lower()
    
    if choice in ['y', 'yes', 'có', '']:
        # Get document name
        print("Tên tài liệu (mặc định: merged_document): ", end='')
        doc_name = input().strip() or "merged_document"
        
        # Create Word and Markdown
        create_merged_document(results, output_dir, doc_name)
        
        # Ask for PDF options
        print("\nTạo PDF? (image/text/both/no): ", end='')
        pdf_choice = input().strip().lower()
        
        if pdf_choice in ['image', 'i', 'both', 'b']:
            # PDF from original images
            image_paths = [r['path'] for r in results]
            pdf_img = output_dir / f"{doc_name}_images.pdf"
            create_pdf_from_images(image_paths, pdf_img)
        
        if pdf_choice in ['text', 't', 'both', 'b']:
            # PDF from OCR text
            pdf_text = output_dir / f"{doc_name}_text.pdf"
            create_pdf_from_ocr(results, pdf_text)
        
        # Ask for Excel export (PHASE 2 FEATURE - for invoices/tables)
        has_any_tables = any(r.get('has_tables', False) for r in results)
        if has_any_tables:
            print("\n💡 Phát hiện bảng biểu! Xuất ra Excel? (y/n): ", end='')
            excel_choice = input().strip().lower()
            
            if excel_choice in ['y', 'yes', 'có', '']:
                excel_file = output_dir / f"{doc_name}_tables.xlsx"
                extract_tables_to_excel(results, excel_file)
        
        # Ask for EPUB
        print("\nTạo EPUB ebook? (y/n): ", end='')
        epub_choice = input().strip().lower()
        
        if epub_choice in ['y', 'yes', 'có', '']:
            print("Tên sách (mặc định: OCR Document): ", end='')
            book_title = input().strip() or "OCR Document"
            print("Tác giả (mặc định: QSM OCR): ", end='')
            author = input().strip() or "QSM OCR"
            
            epub_file = output_dir / f"{doc_name}.epub"
            create_epub(results, epub_file, book_title, author)
    
    # Save individual files
    print(f"\nLưu từng file riêng lẻ? (y/n): ", end='')
    save_individual = input().strip().lower()
    
    if save_individual in ['y', 'yes', 'có', '']:
        print(f"\n[*] Đang lưu từng file...")
        for result in results:
            p = result['path']
            txt = result['text']
            
            # Individual Markdown
            md = output_dir / f"{p.stem}_page.md"
            md.write_text(f"# {p.stem}\n\n{txt}", encoding="utf-8")
            
            # Individual Word
            doc = Document()
            doc.add_heading(f"Trang - {p.stem}", 0)
            doc.add_paragraph(f"Nguồn: {p.name}").italic = True
            for par in txt.split("\n\n"):
                if par.strip():
                    doc.add_paragraph(par.strip())
            dx = output_dir / f"{p.stem}_page.docx"
            doc.save(str(dx))
        
        print(f"    [✓] {len(results)} file đã lưu")
    
    # Done
    print(f"\n{'='*70}")
    print(f"✓ HOÀN TẤT!")
    print(f"Thư mục output: {output_dir}")
    print(f"{'='*70}\n")
    input("Nhấn Enter để thoát...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n[!] Cancelled by user")
        input("Press Enter to exit...")
    except Exception as e:
        print(f"\n[✗] FATAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")
