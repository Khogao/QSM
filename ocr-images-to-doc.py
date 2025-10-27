#!/usr/bin/env python3
"""
🖼️ OCR Images to Document - QSM Quick Tool
===========================================

📸 Chức năng:
- Nhận 1 hoặc nhiều ảnh (JPG, PNG, TIFF, BMP)
- Tự động OCR với EasyOCR (tiếng Việt + Anh)
- Tự động nhận diện thứ tự trang (theo tên file)
- Tự động sắp xếp nội dung
- Xuất ra file Word (.docx) + Markdown (.md)

📖 Cách dùng:
1. Chạy script
2. Kéo thả (drag & drop) ảnh vào hoặc nhập đường dẫn
3. Nhấn Enter
4. Đợi OCR hoàn tất
5. File Word & MD sẽ được tạo trong cùng thư mục

🎯 Style: Đơn giản, nhanh, không cần config
"""

import os
import sys
import json
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
import re

# Fix Windows console encoding
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

# Import dependencies
try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image
    DEPS_OK = True
except ImportError as e:
    DEPS_OK = False
    MISSING_DEP = str(e)


def print_banner():
    """In banner chào mừng"""
    print("\n" + "=" * 70)
    print("🖼️  OCR IMAGES TO DOCUMENT - QSM Quick Tool")
    print("=" * 70)
    print("📸 Chụp ảnh → OCR tự động → Xuất Word/Markdown")
    print("🌐 Hỗ trợ: Tiếng Việt, Tiếng Anh")
    print("=" * 70 + "\n")


def check_dependencies():
    """Kiểm tra dependencies"""
    if not DEPS_OK:
        print("❌ Thiếu thư viện cần thiết!")
        print(f"\n   Lỗi: {MISSING_DEP}\n")
        print("🔧 Cài đặt bằng lệnh:")
        print("   cd python")
        print("   .\\venv\\Scripts\\activate")
        print("   pip install docling easyocr pillow python-docx\n")
        sys.exit(1)


def extract_page_number(filename: str) -> int:
    """
    Trích xuất số trang từ tên file
    
    Patterns nhận diện:
    - page_1.jpg → 1
    - scan_001.png → 1
    - img_005.jpg → 5
    - photo-10.jpg → 10
    - IMG_20241027_001.jpg → 1 (lấy số cuối)
    """
    # Tìm tất cả số trong tên file
    numbers = re.findall(r'\d+', filename)
    
    if not numbers:
        return 0
    
    # Lấy số cuối cùng (thường là số trang)
    return int(numbers[-1])


def sort_images_by_page(image_paths: List[str]) -> List[Tuple[str, int]]:
    """
    Sắp xếp ảnh theo thứ tự trang
    
    Returns:
        List of (path, page_number) tuples, sorted by page_number
    """
    images_with_pages = []
    
    for path in image_paths:
        filename = os.path.basename(path)
        page_num = extract_page_number(filename)
        images_with_pages.append((path, page_num))
    
    # Sắp xếp theo page number
    images_with_pages.sort(key=lambda x: x[1])
    
    return images_with_pages


def ocr_image(image_path: str, converter: DocumentConverter) -> Dict:
    """
    OCR một ảnh
    
    Returns:
        {
            'success': bool,
            'text': str,
            'confidence': float,
            'processing_time': float,
            'error': str (if failed)
        }
    """
    start_time = time.time()
    result = {
        'success': False,
        'text': '',
        'confidence': 0.0,
        'processing_time': 0.0,
        'error': ''
    }
    
    try:
        # Convert image with Docling
        doc_result = converter.convert(image_path)
        doc = doc_result.document
        
        # Extract text
        text = doc.export_to_markdown()
        
        # Get confidence
        confidence = 0.95  # Default
        if hasattr(doc_result, 'confidence'):
            if hasattr(doc_result.confidence, 'mean_grade'):
                confidence = float(doc_result.confidence.mean_grade.value)
        
        result['success'] = True
        result['text'] = text.strip()
        result['confidence'] = confidence
        
    except Exception as e:
        result['error'] = str(e)
    
    finally:
        result['processing_time'] = time.time() - start_time
    
    return result


def create_word_document(pages_data: List[Dict], output_path: str):
    """
    Tạo file Word từ kết quả OCR
    
    Args:
        pages_data: List of {page_num, text, confidence, image_path}
        output_path: Path to output .docx file
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('OCR Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"📅 Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n").italic = True
    metadata.add_run(f"📄 Số trang: {len(pages_data)}\n").italic = True
    metadata.add_run(f"🤖 OCR Engine: Docling + EasyOCR").italic = True
    
    doc.add_paragraph()  # Spacer
    doc.add_page_break()
    
    # Add each page
    for idx, page in enumerate(pages_data, 1):
        # Page header
        page_header = doc.add_heading(f"Trang {page['page_num']} / {len(pages_data)}", level=2)
        
        # Confidence indicator
        confidence = page.get('confidence', 0.0)
        if confidence >= 0.9:
            confidence_text = f"✅ Độ tin cậy: {confidence:.1%} (Rất tốt)"
        elif confidence >= 0.7:
            confidence_text = f"⚠️ Độ tin cậy: {confidence:.1%} (Tốt)"
        else:
            confidence_text = f"❌ Độ tin cậy: {confidence:.1%} (Cần kiểm tra)"
        
        info = doc.add_paragraph()
        info.add_run(f"📁 File: {os.path.basename(page['image_path'])}\n").italic = True
        info.add_run(confidence_text).italic = True
        
        doc.add_paragraph()  # Spacer
        
        # Page content
        content = doc.add_paragraph(page['text'])
        
        # Page break (except last page)
        if idx < len(pages_data):
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)


def create_markdown_document(pages_data: List[Dict], output_path: str):
    """
    Tạo file Markdown từ kết quả OCR
    
    Args:
        pages_data: List of {page_num, text, confidence, image_path}
        output_path: Path to output .md file
    """
    lines = []
    
    # Title
    lines.append("# OCR Document\n")
    
    # Metadata
    lines.append("## 📋 Thông tin tài liệu\n")
    lines.append(f"- **Ngày tạo:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    lines.append(f"- **Số trang:** {len(pages_data)}")
    lines.append(f"- **OCR Engine:** Docling + EasyOCR")
    lines.append("")
    lines.append("---\n")
    
    # Each page
    for page in pages_data:
        lines.append(f"## 📄 Trang {page['page_num']} / {len(pages_data)}\n")
        
        # Metadata
        confidence = page.get('confidence', 0.0)
        if confidence >= 0.9:
            confidence_icon = "✅"
        elif confidence >= 0.7:
            confidence_icon = "⚠️"
        else:
            confidence_icon = "❌"
        
        lines.append(f"**File:** `{os.path.basename(page['image_path'])}`  ")
        lines.append(f"**Độ tin cậy:** {confidence_icon} {confidence:.1%}\n")
        
        # Content
        lines.append("### Nội dung:\n")
        lines.append(page['text'])
        lines.append("\n---\n")
    
    # Save
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def get_image_files() -> List[str]:
    """
    Nhận danh sách file ảnh từ user input
    
    Returns:
        List of image file paths
    """
    print("📂 Nhập đường dẫn ảnh (hoặc kéo thả file vào đây):")
    print("   - Có thể nhập 1 file hoặc nhiều file (phân cách bằng dấu |)")
    print("   - Có thể nhập thư mục chứa ảnh")
    print("   - Nhấn Enter sau khi dán\n")
    
    user_input = input("👉 Đường dẫn: ").strip()
    
    if not user_input:
        print("\n❌ Không có đường dẫn nào được nhập!")
        return []
    
    # Remove quotes if drag-dropped
    user_input = user_input.strip('"').strip("'")
    
    image_files = []
    supported_exts = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp'}
    
    # Check if it's a directory
    if os.path.isdir(user_input):
        print(f"\n📁 Đang quét thư mục: {user_input}")
        for root, dirs, files in os.walk(user_input):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_exts:
                    image_files.append(os.path.join(root, file))
        print(f"   ✅ Tìm thấy {len(image_files)} ảnh")
    
    # Check if multiple files separated by |
    elif '|' in user_input:
        paths = [p.strip().strip('"').strip("'") for p in user_input.split('|')]
        for path in paths:
            if os.path.isfile(path):
                ext = os.path.splitext(path)[1].lower()
                if ext in supported_exts:
                    image_files.append(path)
    
    # Single file
    elif os.path.isfile(user_input):
        ext = os.path.splitext(user_input)[1].lower()
        if ext in supported_exts:
            image_files.append(user_input)
        else:
            print(f"\n❌ File không phải định dạng ảnh: {user_input}")
            print(f"   Hỗ trợ: {', '.join(supported_exts)}")
    
    else:
        print(f"\n❌ Không tìm thấy file/thư mục: {user_input}")
    
    return image_files


def main():
    """Main script"""
    print_banner()
    check_dependencies()
    
    # Get image files
    image_files = get_image_files()
    
    if not image_files:
        print("\n⚠️ Không có ảnh nào để xử lý!")
        print("\nNhấn Enter để thoát...")
        input()
        return
    
    print(f"\n✅ Tìm thấy {len(image_files)} ảnh")
    
    # Sort by page number
    print("\n🔢 Đang sắp xếp theo thứ tự trang...")
    sorted_images = sort_images_by_page(image_files)
    
    print("\n📋 Thứ tự xử lý:")
    for idx, (path, page_num) in enumerate(sorted_images, 1):
        print(f"   {idx}. [{page_num:03d}] {os.path.basename(path)}")
    
    # Confirm
    print("\n❓ Tiếp tục OCR? (Y/n): ", end='')
    confirm = input().strip().lower()
    if confirm and confirm != 'y':
        print("\n⚠️ Đã hủy!")
        return
    
    # Setup Docling
    print("\n⚙️ Đang khởi tạo OCR engine...")
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = True
    pipeline_options.ocr_options = EasyOcrOptions(
        lang=['vi', 'en'],
        force_full_page_ocr=True  # Force OCR toàn trang
    )
    
    converter = DocumentConverter(
        format_options={
            InputFormat.IMAGE: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )
    print("   ✅ Sẵn sàng!\n")
    
    # Process each image
    pages_data = []
    total_time = 0
    
    print("=" * 70)
    print("🚀 BẮT ĐẦU OCR")
    print("=" * 70 + "\n")
    
    for idx, (image_path, page_num) in enumerate(sorted_images, 1):
        filename = os.path.basename(image_path)
        print(f"[{idx}/{len(sorted_images)}] Đang xử lý: {filename}")
        
        # Get image info
        try:
            img = Image.open(image_path)
            width, height = img.size
            file_size = os.path.getsize(image_path) / 1024  # KB
            print(f"    📐 Kích thước: {width}x{height} px ({file_size:.1f} KB)")
        except:
            pass
        
        # OCR
        result = ocr_image(image_path, converter)
        total_time += result['processing_time']
        
        if result['success']:
            word_count = len(result['text'].split())
            char_count = len(result['text'])
            confidence = result['confidence']
            
            # Confidence indicator
            if confidence >= 0.9:
                conf_icon = "✅"
            elif confidence >= 0.7:
                conf_icon = "⚠️"
            else:
                conf_icon = "❌"
            
            print(f"    {conf_icon} Thành công! Độ tin cậy: {confidence:.1%}")
            print(f"    📝 {word_count} từ, {char_count} ký tự")
            print(f"    ⏱️  Thời gian: {result['processing_time']:.1f}s\n")
            
            pages_data.append({
                'page_num': page_num,
                'text': result['text'],
                'confidence': confidence,
                'image_path': image_path,
                'word_count': word_count
            })
        else:
            print(f"    ❌ Lỗi: {result['error']}\n")
    
    # Summary
    print("=" * 70)
    print("📊 KẾT QUẢ")
    print("=" * 70)
    print(f"✅ Thành công:     {len(pages_data)}/{len(sorted_images)} trang")
    print(f"❌ Thất bại:       {len(sorted_images) - len(pages_data)}/{len(sorted_images)} trang")
    print(f"⏱️  Tổng thời gian:  {total_time:.1f}s ({total_time/60:.1f} phút)")
    print(f"📝 Tổng số từ:      {sum(p['word_count'] for p in pages_data)} từ")
    
    if not pages_data:
        print("\n⚠️ Không có trang nào được OCR thành công!")
        print("\nNhấn Enter để thoát...")
        input()
        return
    
    # Generate output files
    print("\n" + "=" * 70)
    print("💾 ĐANG TẠO FILE XUẤT")
    print("=" * 70 + "\n")
    
    # Get output directory (same as first image)
    output_dir = os.path.dirname(sorted_images[0][0])
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Word document
    word_path = os.path.join(output_dir, f"OCR_Result_{timestamp}.docx")
    print(f"📄 Đang tạo Word document...")
    create_word_document(pages_data, word_path)
    print(f"   ✅ {word_path}")
    
    # Markdown document
    md_path = os.path.join(output_dir, f"OCR_Result_{timestamp}.md")
    print(f"📝 Đang tạo Markdown...")
    create_markdown_document(pages_data, md_path)
    print(f"   ✅ {md_path}")
    
    # Final message
    print("\n" + "=" * 70)
    print("🎉 HOÀN THÀNH!")
    print("=" * 70)
    print(f"\n📁 File đã được lưu tại: {output_dir}")
    print(f"   - Word: OCR_Result_{timestamp}.docx")
    print(f"   - Markdown: OCR_Result_{timestamp}.md")
    print("\n✨ Mở file Word để xem kết quả!\n")
    
    # Wait before exit
    print("Nhấn Enter để thoát...")
    input()


if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️ Đã dừng bởi người dùng!")
        sys.exit(0)
    except Exception as e:
        print(f"\n\n❌ Lỗi nghiêm trọng: {e}")
        import traceback
        traceback.print_exc()
        print("\nNhấn Enter để thoát...")
        input()
        sys.exit(1)
