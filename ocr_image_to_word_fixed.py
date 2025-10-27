"""
QSM - OCR Image to Word/Markdown
Quick script: Drag image files -> Auto OCR -> Output Word/Markdown

Usage:
    python ocr_image_to_word.py
    (then paste file paths or drag files)
"""

import sys
import os
import io

# Fix Windows console encoding
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from pathlib import Path
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time

# Supported image formats
SUPPORTED_FORMATS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}

def print_header():
    """Print script header"""
    print("=" * 70)
    print("QSM - OCR Image to Word/Markdown Converter")
    print("=" * 70)
    print("Supported: JPG, PNG, BMP, TIFF")
    print("Languages: Vietnamese + English")
    print("Output: Word (.docx) + Markdown (.md)")
    print("=" * 70)
    print()

def setup_converter():
    """Setup Docling converter with OCR"""
    print("Setting up OCR engine...")
    
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = True
    pipeline_options.do_table_structure = True
    
    # OCR for Vietnamese and English
    pipeline_options.ocr_options = EasyOcrOptions(
        lang=['vi', 'en'],
        force_full_page_ocr=True  # Full page OCR for images
    )
    
    converter = DocumentConverter(
        format_options={
            InputFormat.IMAGE: PdfFormatOption(
                pipeline_options=pipeline_options
            )
        }
    )
    
    print("OCR engine ready!\n")
    return converter

def process_image(image_path, converter, output_dir):
    """Process single image file"""
    try:
        image_path = Path(image_path).resolve()
        
        if not image_path.exists():
            print(f" File not found: {image_path}")
            return False
        
        if image_path.suffix.lower() not in SUPPORTED_FORMATS:
            print(f"  Unsupported format: {image_path.suffix}")
            print(f"   Supported: {', '.join(SUPPORTED_FORMATS)}")
            return False
        
        print(f"\n Processing: {image_path.name}")
        print(f"   Size: {image_path.stat().st_size / 1024:.1f} KB")
        
        # OCR the image
        print("    Running OCR...")
        start_time = time.time()
        
        result = converter.convert(str(image_path))
        doc = result.document
        
        elapsed = time.time() - start_time
        print(f"     OCR completed in {elapsed:.1f}s")
        
        # Get text content
        markdown_text = doc.export_to_markdown()
        
        if not markdown_text.strip():
            print("     No text detected in image")
            return False
        
        # Word count
        word_count = len(markdown_text.split())
        print(f"    Extracted: {word_count} words")
        
        # Save Markdown
        base_name = image_path.stem
        md_file = output_dir / f"{base_name}_ocr.md"
        
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(f"# {base_name}\n\n")
            f.write(f"*OCR extracted from: {image_path.name}*\n\n")
            f.write("---\n\n")
            f.write(markdown_text)
        
        print(f"    Markdown saved: {md_file.name}")
        
        # Save Word document
        docx_file = output_dir / f"{base_name}_ocr.docx"
        create_word_document(markdown_text, docx_file, image_path.name)
        print(f"    Word saved: {docx_file.name}")
        
        # Show confidence if available
        if hasattr(result, 'confidence'):
            try:
                confidence = float(result.confidence.mean_grade.value)
                print(f"    Confidence: {confidence:.1%}")
            except:
                pass
        
        print(f"    SUCCESS!")
        return True
        
    except Exception as e:
        print(f"    Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def create_word_document(text, output_path, source_filename):
    """Create formatted Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('OCR Extracted Content', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Source: {source_filename}\n").italic = True
    meta.add_run(f"Extracted: {time.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"Tool: QSM - QueryMaster OCR").italic = True
    
    # Separator
    doc.add_paragraph("_" * 50)
    
    # Content
    doc.add_heading('Content', 1)
    
    # Split by paragraphs and add to document
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Check if heading
            if para_text.startswith('#'):
                level = para_text.count('#', 0, 3)
                text_clean = para_text.lstrip('#').strip()
                doc.add_heading(text_clean, min(level, 3))
            else:
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.line_spacing = 1.5
    
    doc.save(str(output_path))

def get_image_files():
    """Get image files from user input"""
    print(" Enter image file path(s):")
    print("    Type or paste file path")
    print("    Drag & drop file into terminal")
    print("    Multiple files: separate by newlines")
    print("    Type 'done' when finished")
    print()
    
    files = []
    
    while True:
        user_input = input(">>> ").strip()
        
        if not user_input:
            continue
        
        if user_input.lower() in ['done', 'exit', 'quit', 'q']:
            break
        
        # Clean up path (remove quotes)
        user_input = user_input.strip('"').strip("'")
        
        # Check if file exists
        path = Path(user_input)
        if path.exists() and path.is_file():
            ext = path.suffix.lower()
            if ext in SUPPORTED_FORMATS:
                files.append(path)
                print(f"    Added: {path.name}")
            else:
                print(f"     Unsupported format: {ext}")
        else:
            print(f"    File not found: {user_input}")
    
    return files

def main():
    """Main function"""
    print_header()
    
    # Setup converter
    converter = setup_converter()
    
    # Get image files
    image_files = get_image_files()
    
    if not image_files:
        print("\n  No valid image files provided. Exiting.")
        return
    
    print(f"\n Processing {len(image_files)} file(s)...")
    print("=" * 70)
    
    # Create output directory
    output_dir = Path(__file__).parent / "ocr_output"
    output_dir.mkdir(exist_ok=True)
    
    # Process each file
    success_count = 0
    for img_file in image_files:
        if process_image(img_file, converter, output_dir):
            success_count += 1
    
    # Summary
    print("\n" + "=" * 70)
    print(" SUMMARY")
    print("=" * 70)
    print(f"Total files: {len(image_files)}")
    print(f"Success: {success_count}")
    print(f"Failed: {len(image_files) - success_count}")
    print(f"\n Output folder: {output_dir.absolute()}")
    print("=" * 70)
    print("\n Done! Press Enter to exit...")
    input()

if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n  Interrupted by user. Exiting...")
        sys.exit(0)
    except Exception as e:
        print(f"\n Fatal error: {e}")
        import traceback
        traceback.print_exc()
        print("\nPress Enter to exit...")
        input()
        sys.exit(1)
