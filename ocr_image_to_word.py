import sys, io
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from pathlib import Path
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
from pypdf import PdfWriter
from ebooklib import epub
import time
import re
import uuid

FORMATS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff"}

def header():
    print("="*70)
    print("QSM - OCR Image to Word")
    print("="*70)

def setup():
    print("Setting up OCR...")
    opts = PdfPipelineOptions()
    opts.do_ocr = True
    opts.do_table_structure = True
    opts.ocr_options = EasyOcrOptions(lang=["vi","en"], force_full_page_ocr=True)
    c = DocumentConverter(format_options={InputFormat.IMAGE: PdfFormatOption(pipeline_options=opts)})
    print("Ready!\n")
    return c

def proc(p, c, o):
    """Process single image and return OCR result"""
    try:
        p = Path(p).resolve()
        if not p.exists(): 
            print(f"[X] Not found: {p}")
            return None
        if p.suffix.lower() not in FORMATS:
            print(f"[!] Unsupported: {p.suffix}")
            return None
        print(f"\n[*] {p.name} ({p.stat().st_size/1024:.1f} KB)")
        print("   OCR...")
        t = time.time()
        r = c.convert(str(p))
        print(f"   Done in {time.time()-t:.1f}s")
        txt = r.document.export_to_markdown()
        if not txt.strip():
            print("   [!] No text")
            return None
        print(f"   {len(txt.split())} words")
        
        # Return result dict instead of saving immediately
        return {
            'path': p,
            'text': txt,
            'word_count': len(txt.split()),
            'time': time.time() - t
        }
    except Exception as e:
        print(f"   [ERROR] {e}")
        return None

def extract_page_number(filename):
    """Extract page number from filename for sorting"""
    # Try to find numbers in filename
    numbers = re.findall(r'\d+', filename)
    if numbers:
        # Return first number found (usually page number)
        return int(numbers[0])
    # If no number, use filename for sorting
    return 0

def sort_files_by_page(files):
    """Sort files by page number extracted from filename"""
    return sorted(files, key=lambda f: (extract_page_number(f.stem), f.name))

def create_merged_document(results, output_dir, doc_name="merged_document"):
    """Create merged Word and Markdown from OCR results"""
    
    # Sort results by page number
    sorted_results = sorted(results, key=lambda r: extract_page_number(r['path'].stem))
    
    print("\n[*] Creating merged document...")
    print(f"    Page order: {', '.join([r['path'].stem[:20] + '...' if len(r['path'].stem) > 20 else r['path'].stem for r in sorted_results])}")
    
    # Create merged Markdown
    md_content = f"# Merged OCR Document\n\n"
    md_content += f"*Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
    md_content += f"*Total pages: {len(sorted_results)}*\n\n"
    md_content += "---\n\n"
    
    for idx, result in enumerate(sorted_results, 1):
        md_content += f"## Page {idx} - {result['path'].name}\n\n"
        md_content += result['text']
        md_content += f"\n\n---\n\n"
    
    md_file = output_dir / f"{doc_name}.md"
    md_file.write_text(md_content, encoding="utf-8")
    print(f"    [OK] Markdown: {md_file.name}")
    
    # Create merged Word document
    doc = Document()
    
    # Title page
    title = doc.add_heading('OCR Merged Document', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"Total Pages: {len(sorted_results)}\n").italic = True
    meta.add_run(f"Total Words: {sum(r['word_count'] for r in sorted_results)}\n").italic = True
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Add each page
    for idx, result in enumerate(sorted_results, 1):
        # Page header
        page_heading = doc.add_heading(f'Page {idx}', 1)
        
        # Source info
        source = doc.add_paragraph()
        source.add_run(f"Source: {result['path'].name}\n").italic = True
        source.add_run(f"Words: {result['word_count']}").italic = True
        
        # Add separator
        doc.add_paragraph("_" * 50)
        
        # Content
        paragraphs = result['text'].split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                if para_text.startswith('#'):
                    level = min(para_text.count('#', 0, 3), 2) + 1
                    text_clean = para_text.lstrip('#').strip()
                    doc.add_heading(text_clean, level)
                else:
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.line_spacing = 1.5
        
        # Page break (except last page)
        if idx < len(sorted_results):
            doc.add_page_break()
    
    docx_file = output_dir / f"{doc_name}.docx"
    doc.save(str(docx_file))
    print(f"    [OK] Word: {docx_file.name}")
    
    return md_file, docx_file

def create_pdf_from_images(image_paths, output_path):
    """Create PDF from original images (preserving quality)"""
    try:
        print(f"\n[*] Creating PDF from original images...")
        
        # Open all images and convert to RGB
        images = []
        for img_path in image_paths:
            img = Image.open(img_path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)
        
        # Save as PDF
        if images:
            images[0].save(
                output_path,
                save_all=True,
                append_images=images[1:],
                resolution=100.0,
                quality=95,
                optimize=True
            )
            print(f"    [OK] Image PDF: {output_path.name} ({len(images)} pages)")
            return output_path
    except Exception as e:
        print(f"    [ERROR] Failed to create image PDF: {e}")
        return None

def create_pdf_from_ocr(results, output_path):
    """Create PDF from OCR text content"""
    try:
        print(f"\n[*] Creating PDF from OCR text...")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2C3E50',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#34495E',
            spaceAfter=12,
            spaceBefore=12
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=10
        )
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=9,
            textColor='#7F8C8D',
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        # Build content
        story = []
        
        # Title page
        story.append(Paragraph("OCR Merged Document", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
        story.append(Paragraph(f"Total Pages: {len(results)}", meta_style))
        story.append(Paragraph(f"Total Words: {sum(r['word_count'] for r in results)}", meta_style))
        story.append(PageBreak())
        
        # Add each page
        for idx, result in enumerate(results, 1):
            # Page header
            story.append(Paragraph(f"Page {idx}", heading_style))
            story.append(Paragraph(f"Source: {result['path'].name}", meta_style))
            story.append(Spacer(1, 0.1*inch))
            
            # Content
            text = result['text']
            paragraphs = text.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Clean text for PDF
                    clean_text = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    # Check if heading
                    if clean_text.startswith('#'):
                        clean_text = clean_text.lstrip('#').strip()
                        story.append(Paragraph(clean_text, heading_style))
                    else:
                        story.append(Paragraph(clean_text, normal_style))
            
            # Page break (except last)
            if idx < len(results):
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        print(f"    [OK] Text PDF: {output_path.name}")
        return output_path
        
    except Exception as e:
        print(f"    [ERROR] Failed to create text PDF: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_epub(results, output_path, book_title="OCR Document", author="QSM OCR"):
    """Create EPUB ebook from OCR results"""
    try:
        print(f"\n[*] Creating EPUB ebook...")
        
        # Create book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(str(uuid.uuid4()))
        book.set_title(book_title)
        book.set_language('vi')  # Vietnamese
        book.add_author(author)
        
        # Create chapters
        chapters = []
        spine = ['nav']
        
        for idx, result in enumerate(results, 1):
            # Create chapter
            chapter = epub.EpubHtml(
                title=f'Page {idx}',
                file_name=f'page_{idx}.xhtml',
                lang='vi'
            )
            
            # Build HTML content
            content = f'<h1>Page {idx}</h1>'
            content += f'<p><em>Source: {result["path"].name}</em></p>'
            content += '<hr/>'
            
            # Convert markdown to HTML
            text = result['text']
            paragraphs = text.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Clean HTML
                    clean = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    if clean.startswith('#'):
                        # Heading
                        level = min(clean.count('#', 0, 3), 3)
                        text_clean = clean.lstrip('#').strip()
                        content += f'<h{level + 1}>{text_clean}</h{level + 1}>'
                    else:
                        # Paragraph
                        content += f'<p>{clean}</p>'
            
            chapter.content = content
            
            # Add to book
            book.add_item(chapter)
            chapters.append(chapter)
            spine.append(chapter)
        
        # Add default NCX and Nav files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = '''
            @namespace epub "http://www.idpf.org/2007/ops";
            body { font-family: Georgia, serif; margin: 2em; }
            h1 { color: #2C3E50; border-bottom: 2px solid #3498DB; padding-bottom: 0.3em; }
            h2, h3 { color: #34495E; }
            p { text-align: justify; line-height: 1.6; }
            em { color: #7F8C8D; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Create spine
        book.spine = spine
        
        # Add table of contents
        book.toc = tuple(chapters)
        
        # Write EPUB
        epub.write_epub(str(output_path), book, {})
        print(f"    [OK] EPUB: {output_path.name} ({len(chapters)} chapters)")
        return output_path
        
    except Exception as e:
        print(f"    [ERROR] Failed to create EPUB: {e}")
        import traceback
        traceback.print_exc()
        return None

def get_files():
    print("Enter file paths (drag & drop or type, then 'done'):\n")
    fs = []
    while True:
        i = input(">>> ").strip()
        if not i: continue
        if i.lower() in ["done","q","exit"]: break
        
        # Split by quotes for multiple files (Windows drag & drop)
        import re
        paths = re.findall(r'"([^"]+)"', i)
        if not paths:
            # Single file without quotes
            paths = [i.strip('"').strip("'")]
        
        for path_str in paths:
            p = Path(path_str.strip())
            if p.exists() and p.suffix.lower() in FORMATS:
                fs.append(p)
                print(f"   [+] {p.name}")
            else:
                print(f"   [X] Invalid: {p.name if p.name else path_str}")
    return fs

def main():
    header()
    c = setup()
    fs = get_files()
    if not fs:
        print("\n[!] No files")
        return
    
    # Sort files by page number
    fs = sort_files_by_page(fs)
    
    print(f"\nProcessing {len(fs)} file(s) in order...")
    print("="*70)
    
    o = Path(__file__).parent / "ocr_output"
    o.mkdir(exist_ok=True)
    
    # Process all files and collect results
    results = []
    for f in fs:
        result = proc(f, c, o)
        if result:
            results.append(result)
    
    # Create merged document if we have results
    if results:
        print(f"\n{'='*70}")
        print(f"Successfully OCR'd: {len(results)}/{len(fs)} files")
        
        # Ask if user wants merged document
        print("\nCreate merged document? (y/n): ", end='')
        choice = input().strip().lower()
        
        if choice in ['y', 'yes', '']:
            # Get document name
            print("Enter document name (default: merged_document): ", end='')
            doc_name = input().strip() or "merged_document"
            
            # Create Word and Markdown
            create_merged_document(results, o, doc_name)
            
            # Ask for PDF options
            print("\nCreate PDF? (image/text/both/no): ", end='')
            pdf_choice = input().strip().lower()
            
            if pdf_choice in ['image', 'i', 'both', 'b']:
                # PDF from original images
                image_paths = [r['path'] for r in results]
                pdf_img = o / f"{doc_name}_images.pdf"
                create_pdf_from_images(image_paths, pdf_img)
            
            if pdf_choice in ['text', 't', 'both', 'b']:
                # PDF from OCR text
                pdf_text = o / f"{doc_name}_text.pdf"
                create_pdf_from_ocr(results, pdf_text)
            
            # Ask for EPUB
            print("\nCreate EPUB ebook? (y/n): ", end='')
            epub_choice = input().strip().lower()
            
            if epub_choice in ['y', 'yes', '']:
                print("Book title (default: OCR Document): ", end='')
                book_title = input().strip() or "OCR Document"
                print("Author (default: QSM OCR): ", end='')
                author = input().strip() or "QSM OCR"
                
                epub_file = o / f"{doc_name}.epub"
                create_epub(results, epub_file, book_title, author)
        
        # Also save individual files
        print(f"\n[*] Saving individual files...")
        for result in results:
            p = result['path']
            txt = result['text']
            
            # Individual Markdown
            md = o / f"{p.stem}_page.md"
            md.write_text(f"# {p.stem}\n\n{txt}", encoding="utf-8")
            
            # Individual Word
            doc = Document()
            doc.add_heading(f"Page - {p.stem}", 0)
            doc.add_paragraph(f"Source: {p.name}").italic = True
            for par in txt.split("\n\n"):
                if par.strip():
                    doc.add_paragraph(par.strip())
            dx = o / f"{p.stem}_page.docx"
            doc.save(str(dx))
        
        print(f"    [OK] {len(results)} individual files saved")
    
    print(f"\n{'='*70}\nOutput: {o}\n{'='*70}")
    input("Press Enter...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[!] Stopped")
    except Exception as e:
        print(f"\n[ERROR] {e}")
        input()
